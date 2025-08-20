# parsers/docx_parser.py
from typing import List, Optional
from models import DocElement
from pathlib import Path
from docx import Document
from lxml import etree

OMML_NS = {"m": "http://schemas.openxmlformats.org/officeDocument/2006/math"}

def parse_docx(path: str, category: Optional[str]=None) -> List[DocElement]:
    doc = Document(path)
    stem = Path(path).stem
    out: List[DocElement] = []
    doc_id = stem
    order = 0

    # Параграфы и списки
    for para in doc.paragraphs:
        txt = (para.text or "").strip()
        if txt:
            out.append(DocElement(
                category=category, doc_path=path, doc_id=doc_id, source_type="docx",
                element_type="paragraph", text=txt, order=order
            ))
            order += 1

    # Таблицы
    for tbl in doc.tables:
        for row in tbl.rows:
            cells = [ (c.text or "").replace("\n"," ").strip() for c in row.cells ]
            out.append(DocElement(
                category=category, doc_path=path, doc_id=doc_id, source_type="docx",
                element_type="table", text="\t".join(cells), order=order
            ))
            order += 1

    # Формулы OMML безопасно через корневой XML
    try:
        root = doc._element  # CT_Document
        nodes = root.xpath(".//m:oMath", namespaces=OMML_NS)
        for node in nodes:
            xml_str = etree.tostring(node, encoding="unicode")
            out.append(DocElement(
                category=category, doc_path=path, doc_id=doc_id, source_type="docx",
                element_type="formula", html=xml_str, order=order
            ))
            order += 1
    except Exception:
        # молча пропускаем формулы, если структура нестандартная
        pass

    return out